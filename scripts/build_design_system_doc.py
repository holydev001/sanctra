from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "docs/Sanctra_Design_System_v1.docx"
INK = RGBColor(15, 20, 34)
GOLD = RGBColor(183, 122, 0)
MUTED = RGBColor(74, 84, 104)
LIGHT = "F2F4F7"
DARK = "0A0A0A"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
for name, size, before, after in [("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 11.5, 8, 4)]:
    s = styles[name]; s.font.name = "Arial"; s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = GOLD
    s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True

def set_font(run, name="Arial", size=10.5, bold=False, color=INK, italic=False):
    run.font.name = name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name); run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr(); el = tc.first_child_found_in("w:tcMar")
    if el is None: el = OxmlElement("w:tcMar"); tc.append(el)
    for edge, val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        n=OxmlElement(f"w:{edge}"); n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa"); el.append(n)

def table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    trPr = t.rows[0]._tr.get_or_add_trPr(); tblHeader = OxmlElement("w:tblHeader"); tblHeader.set(qn("w:val"), "true"); trPr.append(tblHeader)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]); shade(c,DARK); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r=c.paragraphs[0].add_run(h); set_font(r,size=9,bold=True,color=RGBColor(245,179,36))
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].width=Inches(widths[i]); margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r=cells[i].paragraphs[0].add_run(str(v)); set_font(r,size=9.25)
            if len(t.rows)%2==1: shade(cells[i],LIGHT)
    t.style="Table Grid"
    return t

def bullet(text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25); p.paragraph_format.space_after=Pt(4)
    p.add_run(text)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT
set_font(header.add_run("SANCTRA  /  DESIGN SYSTEM"),size=8.5,bold=True,color=MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("Implementation-derived reference  •  v1.0  •  July 2026"),size=8,color=MUTED)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(4)
set_font(p.add_run("SANCTRA"),size=11,bold=True,color=GOLD)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
set_font(p.add_run("Design System"),size=28,bold=True,color=INK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18)
set_font(p.add_run("A complete implementation-derived guide for consistent public-site and future product design."),size=13,color=MUTED)
table(["Source", "Status", "Authority"], [["sanct-proto TanStack implementation", "Migrated to Next.js", "Exact values in globals.css and design-tokens.ts"],["Document version", "1.0", "Brand and interface reference"]],[2.2,1.5,2.8])

doc.add_heading("1. Brand foundation",1)
doc.add_paragraph("Sanctra is an engineering control plane. Its interface should feel operational, deliberate, trustworthy, and technically literate. The visual system is dark-first, restrained, square, and precise. Gold is used as a controlled signal rather than decoration.")
doc.add_heading("Design principles",2)
for x in ["Operational clarity: visuals explain repository state, infrastructure flow, or action.","Hard geometry: square corners, hairline borders, and disciplined grids.","One signal color: gold identifies brand, focus, and primary action.","Calm density: technical information stays compact without becoming noisy.","Reversible trust: UI language reinforces reviewable pull requests and user control."]: bullet(x)
doc.add_heading("Anti-patterns",2)
doc.add_paragraph("Do not introduce rounded SaaS cards, purple AI gradients, mascots, excessive glassmorphism, fake product metrics, generic stock illustration, playful bounce motion, or a second decorative accent color.")

doc.add_heading("2. Color system",1)
table(["Token", "Dark / default", "Light / opt-in"], [["Background","#050505","#F7F8FB"],["Foreground","#F5F7FA","#0F1422"],["Card","#0A0A0A","#FFFFFF"],["Surface","#0D0D0D","#FFFFFF"],["Surface muted","#171717","#EEF1F6"],["Primary","#F5B324","#B77A00"],["Muted text","#A1A1AA","#4A5468"],["Border","white / 8%","#0F1422 / 10%"],["Input","white / 14%","#0F1422 / 16%"]],[1.8,2.35,2.35])
doc.add_heading("Semantic colors",2)
table(["Meaning", "Value", "Usage"], [["Success","#2EAD78","Passing checks, available integrations"],["Warning","#D99A32","Attention and non-blocking risk"],["Destructive","#D95C5C","Errors and removal"],["Info","#5B8DEF","Neutral system information"]],[1.5,1.2,3.8])
doc.add_paragraph("Accessibility rule: never communicate status through color alone. Pair color with text, a symbol, a dot, or a structural treatment.")

doc.add_heading("3. Typography",1)
table(["Role", "Family", "Specification"], [["Display","Space Grotesk","Headings; 600 weight; -0.02em tracking"],["Body/UI","Inter","400–700; cv11, ss01, ss03 features"],["Code/ops","JetBrains Mono","Logs, YAML, labels, branches, metrics"],["Mono label","JetBrains Mono","11px; 500; 0.12em; uppercase"]],[1.45,1.75,3.3])
doc.add_paragraph("Use fluid clamp-based display sizes for major marketing headlines. Preserve readable line lengths, compact UI labels, and explicit hierarchy. Fonts are loaded through next/font and mapped to CSS variables.")

doc.add_heading("4. Layout and geometry",1)
table(["Property", "Value", "Application"], [["Global radius","0px","All elements and pseudo-elements"],["Content width","1280px","Primary marketing container"],["Page gutter","24px","Mobile and desktop outer padding"],["Section rhythm","80px / 96px","Base / md vertical padding"],["Background grid","48px","Subtle masked hero/section grid"],["Border","1px hairline","Cards, sections, controls"],["Hover lift","3px","Interactive cards only"]],[1.65,1.35,3.5])
doc.add_heading("Surface hierarchy",2)
doc.add_paragraph("Background is the canvas. Card contains discrete modules. Surface represents operational elevation. Surface-muted is reserved for quiet secondary regions. Shadows are rare and must remain restrained and gold-tinted only where interaction benefits from lift.")

doc.add_heading("5. Component specifications",1)
table(["Component", "Visual contract", "Reuse"], [["SiteNav","Sticky translucent canvas, blur, bottom hairline, clipped label swap","All public pages"],["Button","Square, bordered, compact; gold primary; clear focus ring","Global"],["Pill","Mono label, border, tint, optional status dot","Global"],["Card","Hairline + card surface; optional 3px lift","Global"],["Terminal","#0B1120, 12.5px mono, 1.7 line height, diff rails","Marketing/product"],["Form","Explicit labels, square controls, gold focus state","Global"],["Footer","Three colonnades, system status, oversized wordmark","Public"],["SanctraMark","Square S monogram; optional wordmark","Global"]],[1.35,3.75,1.4])
doc.add_heading("Buttons and states",2)
for x in ["Primary: solid gold with dark foreground; use once per decision cluster.","Secondary: bordered surface with foreground text.","Ghost: no visible border until interaction; use for low-priority actions.","Focus: 2px semantic ring with offset against the current background.","Disabled/loading: preserve width and label context; never rely on opacity alone."]: bullet(x)

doc.add_heading("6. Motion language",1)
table(["Pattern", "Timing", "Purpose"], [["Scroll reveal","800ms; cubic-bezier(.2,.7,.2,1); 14px","Section entrance"],["Rise","900ms; same easing; 18px","One-shot mount"],["Scan","3.4s ease-in-out","Repository/process analysis"],["Pulse dot","1.8s ease-in-out","Live status"],["Ticker","40s linear","Continuous trust/technology strip"],["Marquee","60s linear","Slow ambient sequence"],["Type-in","1.2s steps(40)","Code/command disclosure"],["Stamp","700ms","Approval/result emphasis"]],[1.35,2.6,2.55])
doc.add_paragraph("All reveal and looping motion must be removed under prefers-reduced-motion. Avoid layout thrashing, large parallax, and decorative movement that does not communicate hierarchy or state.")

doc.add_heading("7. Responsive and accessibility rules",1)
table(["Breakpoint", "Width", "Typical change"], [["sm","640px","CTA and compact grid expansion"],["md","768px","Desktop navigation; section rhythm; multi-column layouts"],["lg","1024px","Primary hero and feature compositions"],["xl","1280px","Full container composition"],["2xl","1536px","Large viewport breathing room"]],[1.35,1.2,3.95])
for x in ["Semantic heading order and native interactive elements are mandatory.","Keyboard focus must remain clearly visible on every action.","Navigation, FAQ controls, and forms must be fully keyboard operable.","Code panes scroll horizontally rather than shrinking text.","No horizontal page overflow at 360px and above.","Text and status colors must meet practical WCAG 2.2 AA contrast."]: bullet(x)

doc.add_heading("8. Content and visual language",1)
doc.add_paragraph("Copy is clear, technical, calm, and concrete. Prefer “Connect a repository and open a reviewable infrastructure pull request” over broad claims about transforming workflows. Illustrations should be infrastructure visualizations: repositories, pipeline stages, files, checks, diffs, and deployment targets. The product itself is the hero.")
doc.add_heading("Design review checklist",2)
for x in ["Uses only approved colors and font families.","Preserves global square geometry and hairline borders.","Uses gold for signal, not decoration.","Matches 1280px container and 24px gutter conventions.","Includes hover, focus, loading, error, empty, and reduced-motion states.","Uses operational imagery rather than generic AI imagery.","Works at 360, 390, 768, 1024, 1280, and 1440px.","Adds reusable tokens before repeating arbitrary values."]: bullet(x)

doc.add_heading("9. Source of truth",1)
doc.add_paragraph("Machine-readable tokens live in src/config/design-tokens.ts. Runtime CSS variables, utilities, themes, and animations live in src/app/globals.css. Component behavior lives in src/components. This document explains the rules; implementation files remain authoritative when exact code values are required.")
doc.add_paragraph("Change process: propose the token or component-contract change, validate it in both themes and required breakpoints, update code and documentation together, and record any intentional deviation from the original prototype.")

doc.save(OUT)
print(OUT)
